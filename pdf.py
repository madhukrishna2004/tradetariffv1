from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    # Create the hyperlink
    run = paragraph.add_run(text)
    r = run._r
    rPr = r.get_or_add_rPr()
    
    # Add the hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), url)
    
    # Wrap the run inside the hyperlink tag
    hyperlink.append(r)
    paragraph._element.append(hyperlink)

def generate_stylish_welcome_pdf(output_path):
    # Create a new Document
    document = Document()

    # Title
    title = document.add_heading("Welcome to TradeSphere Global", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add introduction with styling
    intro = document.add_paragraph()
    intro_run = intro.add_run(
        "Thank you for subscribing to the Enterprise Plan of TradeSphere Global. "
        "We are excited to have you on board! TradeSphere Global empowers you with "
        "powerful tools for managing global trade tariffs, commodity codes, and rules of origin."
    )
    intro_run.font.size = Pt(12)
    intro_run.font.color.rgb = RGBColor(0, 0, 102)  # Navy Blue
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Steps to access data
    document.add_heading("Steps to Access Data", level=2)
    steps = [
        "1. Visit the official website: ",
        "2. Log in with your credentials provided below.",
    ]
    for step in steps:
        step_paragraph = document.add_paragraph(step)
        step_paragraph.style = "List Bullet"
    
    # Add clickable link to website
    link_paragraph = document.add_paragraph()
    add_hyperlink(link_paragraph, 'https://tradetariffv1.onrender.com', 'https://tradetariffv1.onrender.com')

    # Credentials Section
    document.add_heading("Credentials", level=2)
    credentials_table = document.add_table(rows=1, cols=3)
    credentials_table.style = "Table Grid"
    header_cells = credentials_table.rows[0].cells
    header_cells[0].text = "Role"
    header_cells[1].text = "Username"
    header_cells[2].text = "Password"

    credentials = [
        ("Member 1", "tradesphere@user3", "user3"),
        ("Member 2", "tradesphere@user4", "user4"),
        ("Member 3", "tradesphere@user5", "user5"),
        ("Member 4", "tradesphere@user6", "user6"),
        ("Member 5", "tradesphere@user7", "user7"),
    ]

    for role, username, password in credentials:
        row_cells = credentials_table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = username
        row_cells[2].text = password

    # Add "Explore All Features" Section
    document.add_heading("Explore All Features", level=2)
    explore_paragraph = document.add_paragraph()
    add_hyperlink(explore_paragraph, 'https://tradetariffv1.onrender.com', 'Click here to explore all features of TradeSphere Global')

    # Footer Section
    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "© 2025 TradeSphere Global. All rights reserved. For support, contact tradesphereglobalqueries@gmail.com"
    )
    footer_run.font.size = Pt(10)
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save the Document
    document.save(output_path)
    print(f"Document saved at: {output_path}")


# File paths
output_docx = "welcome_guidelines.docx"
generate_stylish_welcome_pdf(output_docx)
